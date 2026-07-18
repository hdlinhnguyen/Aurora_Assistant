from io import BytesIO

from docx import Document

from test_lifecycle_callbacks import prepared_exam


def test_exports_docx_with_question_answer_rubric_and_topics(client, teacher_headers):
    exam = prepared_exam(client, teacher_headers)
    response = client.post(
        f"/api/exams/{exam['exam_id']}/exports/docx",
        headers=teacher_headers,
        json={
            "style": "standard",
            "include_answer_key": True,
            "include_rubric": True,
            "expected_version": exam["version"],
        },
    )
    assert response.status_code == 201
    export = response.json()
    assert export["exam_version"] == exam["version"]
    assert "file_path" not in export
    listed = client.get(
        f"/api/exams/{exam['exam_id']}/exports", headers=teacher_headers
    )
    assert [item["export_id"] for item in listed.json()] == [export["export_id"]]

    download = client.get(
        f"/api/exams/{exam['exam_id']}/exports/{export['export_id']}/download",
        headers=teacher_headers,
    )
    assert download.status_code == 200
    document = Document(BytesIO(download.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Đề tự luận" in text
    assert "ĐÁP ÁN VÀ BAREM" in text
    assert "topic-linear-equations" in text
    assert "Tổng điểm: 10.00" in text
    audit = client.get(
        f"/api/exams/{exam['exam_id']}/audit", headers=teacher_headers
    ).json()
    assert "docx_exported" in {item["action"] for item in audit}


def test_compact_export_can_omit_answers_and_uses_safe_filename(
    client, teacher_headers
):
    exam = prepared_exam(client, teacher_headers)
    renamed = client.patch(
        f"/api/exams/{exam['exam_id']}",
        headers=teacher_headers,
        json={
            "title": "../../Đề: Toán * số 1",
            "expected_version": exam["version"],
        },
    ).json()
    response = client.post(
        f"/api/exams/{exam['exam_id']}/exports/docx",
        headers=teacher_headers,
        json={
            "style": "compact",
            "include_answer_key": False,
            "include_rubric": False,
            "expected_version": renamed["version"],
        },
    )
    assert response.status_code == 201
    metadata = response.json()
    assert "/" not in metadata["file_name"]
    assert "\\" not in metadata["file_name"]
    download = client.get(
        f"/api/exams/{exam['exam_id']}/exports/{metadata['export_id']}/download",
        headers=teacher_headers,
    )
    document = Document(BytesIO(download.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "ĐÁP ÁN VÀ BAREM" not in text

    returned = client.post(
        f"/api/exams/{exam['exam_id']}/return-to-draft",
        headers=teacher_headers,
        json={"expected_version": renamed["version"]},
    ).json()
    deleted = client.delete(
        f"/api/exams/{exam['exam_id']}",
        headers=teacher_headers,
        params={"expected_version": returned["version"]},
    )
    assert deleted.status_code == 204
    assert (
        client.get(
            f"/api/exams/{exam['exam_id']}/exports/{metadata['export_id']}/download",
            headers=teacher_headers,
        ).status_code
        == 404
    )


def test_docx_renders_single_choice_options_answer_and_owner_scope(
    client, teacher_headers
):
    exam = client.post(
        "/api/exams",
        headers=teacher_headers,
        json={
            "title": "Đề trắc nghiệm",
            "subject_id": "math",
            "grade_level": 8,
            "duration_minutes": 15,
            "total_points": "2.00",
        },
    ).json()
    client.post(
        f"/api/exams/{exam['exam_id']}/questions/from-bank",
        headers=teacher_headers,
        json={
            "question_id": "bank-math-1",
            "points": "2.00",
            "expected_version": 1,
        },
    )
    prepared = client.post(
        f"/api/exams/{exam['exam_id']}/prepare",
        headers=teacher_headers,
        json={"expected_version": 2},
    ).json()
    metadata = client.post(
        f"/api/exams/{exam['exam_id']}/exports/docx",
        headers=teacher_headers,
        json={
            "style": "standard",
            "include_answer_key": True,
            "include_rubric": False,
            "expected_version": prepared["version"],
        },
    ).json()
    download = client.get(
        f"/api/exams/{exam['exam_id']}/exports/{metadata['export_id']}/download",
        headers=teacher_headers,
    )
    document = Document(BytesIO(download.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "A. 1" in text
    assert "B. 2" in text
    assert "Đáp án: b" in text

    other_teacher = {
        "X-Teacher-Id": "teacher-other",
        "X-Role": "teacher",
    }
    assert (
        client.get(
            f"/api/exams/{exam['exam_id']}/exports/{metadata['export_id']}/download",
            headers=other_teacher,
        ).status_code
        == 404
    )
