from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


class DocumentExporter:
    def export(
        self,
        snapshot: dict,
        style: str,
        include_answer_key: bool,
        include_rubric: bool,
        destination: Path,
    ) -> None:
        document = Document()
        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2)
        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(12 if style == "standard" else 10.5)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(snapshot["title"])
        run.bold = True
        run.font.size = Pt(16)
        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(
            f"Môn: {snapshot['subject_id']} · Khối: {snapshot['grade_level']} · "
            f"Thời gian: {snapshot['duration_minutes']} phút · "
            f"Tổng điểm: {snapshot['total_points']}"
        )
        document.add_paragraph("Họ và tên: ____________________  Lớp: ______")
        if snapshot["instructions"]:
            document.add_paragraph(f"Hướng dẫn: {snapshot['instructions']}")

        for index, question in enumerate(snapshot["questions"], start=1):
            paragraph = document.add_paragraph()
            paragraph.add_run(f"Câu {index} ({question['points']} điểm). ").bold = True
            paragraph.add_run(question["content"])
            for choice_index, choice in enumerate(question["choices"]):
                letter = chr(ord("A") + choice_index)
                document.add_paragraph(f"{letter}. {choice['content']}", style=None)
            if question["question_type"] == "essay":
                for _ in range(5 if style == "standard" else 2):
                    document.add_paragraph(
                        "................................................................................"
                    )

        if include_answer_key or include_rubric:
            document.add_page_break()
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.add_run("ĐÁP ÁN VÀ BAREM").bold = True
            for index, question in enumerate(snapshot["questions"], start=1):
                document.add_paragraph(f"Câu {index}", style="Heading 2")
                if include_answer_key and question["question_type"] == "single_choice":
                    document.add_paragraph(f"Đáp án: {question['correct_choice_id']}")
                if include_rubric and question["question_type"] == "essay":
                    for item in question["rubric_items"]:
                        topics = ", ".join(item["topic_ids"])
                        document.add_paragraph(
                            f"- {item['description']} — {item['points']} điểm "
                            f"— Topics: {topics}"
                        )
            document.add_paragraph(f"Tổng điểm: {snapshot['total_points']}")
        destination.parent.mkdir(parents=True, exist_ok=True)
        document.save(destination)
